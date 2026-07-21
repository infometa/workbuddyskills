#!/usr/bin/env python3
"""
票据自动扫描与分类整理脚本 v1.2
扫描指定文件夹中的票据/发票文件（图片、PDF、DOCX、Excel），
自动按类型分类、重命名，并生成含完整发票要素的汇总台账。

⚠️ 重要警示：脚本输出仅供预分类与打草稿。
    金额、日期、发票号任一字段都必须由财务人工逐张复核，
    复核通过才能进入会计系统入账。

v1.2 更新（修复 2026-04-21 回归测试 P0 问题）:
- 修复金额截断 bug：原正则对 "2143.75" 会截成 "214"，现全部重写
- 新增 DOCX 支持：用 python-docx 抽取打印版发票/报销单文本
- 复核清单诚实化：OCR 结果、缺字段的 PDF 一律进"需人工复核"
- 分类归档按票据类型落地：01_增值税/02_电子发票/03_捐赠/04_差旅/05_餐饮/06_其他
- 分类关键词扩展："增值税专票/专用/普票/普通" 都能命中

v1.1 历史:
- PDF 文本抽取 + easyocr 图片识别

用法:
    python receipt_organizer.py <票据文件夹路径> [--output <输出文件夹路径>]
    python receipt_organizer.py <票据文件夹路径> --scan-only

依赖:
    - Python >= 3.8（无需 3.10+ 特性，兼容 3.8~3.14）
    - pdfplumber (PDF 文本提取)     pip install pdfplumber
    - openpyxl (Excel 输出)         pip install openpyxl
    - python-docx (DOCX 支持)       pip install python-docx
    - easyocr + Pillow (图片 OCR)   pip install easyocr pillow   [可选]
"""

import os
import sys
import re
import shutil
import csv
import json
from datetime import datetime
from pathlib import Path
from collections import defaultdict

# 可选依赖
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import docx  # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# OCR 懒加载（首次调用时才初始化，避免纯 PDF 场景加载 torch）
try:
    import easyocr  # noqa: F401
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False

_OCR_READER = None  # 单例


def _get_ocr_reader():
    """首次调用时初始化 easyocr Reader（加载约 3~10 秒）"""
    global _OCR_READER
    if _OCR_READER is not None:
        return _OCR_READER
    if not HAS_EASYOCR:
        return None
    try:
        import easyocr
        print("  [OCR] 首次调用 OCR，正在初始化中英文模型（约 3~10 秒）...")
        _OCR_READER = easyocr.Reader(["ch_sim", "en"], gpu=False, verbose=False)
        return _OCR_READER
    except Exception as e:
        print(f"  [WARN] OCR 初始化失败: {e}")
        return None

# ============================================================
# 配置
# ============================================================

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
EXCEL_EXTS = {".xlsx", ".xls", ".csv"}
ALL_SUPPORTED = IMAGE_EXTS | PDF_EXTS | DOCX_EXTS | EXCEL_EXTS

# 六大类标准归档目录（对齐 SKILL.md 的交付物结构）
CATEGORY_ARCHIVE_DIR = {
    "增值税专用发票": "01_增值税发票",
    "增值税普通发票": "01_增值税发票",
    "电子发票": "02_电子发票",
    "捐赠票据": "03_公益捐赠票据",
    "差旅-火车": "04_差旅报销",
    "差旅-飞机": "04_差旅报销",
    "差旅-出租": "04_差旅报销",
    "差旅-住宿": "04_差旅报销",
    "差旅-其他": "04_差旅报销",
    "餐饮": "05_餐饮",
    "其他": "06_其他",
}

# 分类关键词（按特异性排序，匹配时按长度倒序）
# key 是关键词（可同时出现在文件名或文件内容中）
CATEGORY_KEYWORDS = [
    # 捐赠类（最优先，避免被"发票"误吃）
    ("公益事业捐赠票据", "捐赠票据"),
    ("公益事业捐赠", "捐赠票据"),
    ("捐赠票据", "捐赠票据"),
    ("捐赠收据", "捐赠票据"),
    ("捐赠", "捐赠票据"),
    # 差旅类（细分子类）
    ("航空运输电子客票", "差旅-飞机"),
    ("飞机行程单", "差旅-飞机"),
    ("行程单", "差旅-飞机"),
    ("机票", "差旅-飞机"),
    ("火车票", "差旅-火车"),
    ("高铁", "差旅-火车"),
    ("动车", "差旅-火车"),
    ("出租车", "差旅-出租"),
    ("打车票", "差旅-出租"),
    ("住宿", "差旅-住宿"),
    ("酒店", "差旅-住宿"),
    # 餐饮类
    ("餐饮", "餐饮"),
    ("餐费", "餐饮"),
    # 增值税专/普票（关键词多形式）
    ("增值税专用发票", "增值税专用发票"),
    ("增值税专票", "增值税专用发票"),
    ("增值税普通发票", "增值税普通发票"),
    ("增值税普票", "增值税普通发票"),
    # 电子发票
    ("电子发票", "电子发票"),
    ("电子普通发票", "电子发票"),
    # 兜底（放最后，避免误吃）
    ("发票", "其他"),
    ("报销单", "其他"),
]

# 费用科目建议
EXPENSE_KEYWORDS = {
    "救灾": "业务活动成本—慈善项目支出",
    "物资": "业务活动成本—慈善项目支出",
    "捐赠": "捐赠收入",
    "审计": "管理费用—专业服务费",
    "咨询": "管理费用—专业服务费",
    "法律": "管理费用—专业服务费",
    "网站": "管理费用—专业服务费",
    "活动": "业务活动成本",
    "会议": "业务活动成本/管理费用",
    "差旅": "差旅费（按出差目的归类）",
    "交通": "差旅费（按出差目的归类）",
    "火车": "差旅费（按出差目的归类）",
    "出租": "差旅费（按出差目的归类）",
    "飞机": "差旅费（按出差目的归类）",
    "住宿": "差旅费（按出差目的归类）",
    "餐饮": "业务活动成本—志愿者补贴/会议费",
    "办公": "管理费用—办公费",
    "打印": "管理费用—办公费",
    "房租": "管理费用—房租物业",
    "物业": "管理费用—房租物业",
    "水电": "管理费用—房租物业",
    "电话": "管理费用—通讯费",
    "网络": "管理费用—通讯费",
    "培训": "业务活动成本/管理费用—培训费",
}


# ============================================================
# PDF 发票要素提取（v2 新增）
# ============================================================

# 发票/票据代码/号码正则（同时覆盖"发票"和"票据"两种叫法）
RE_INVOICE_CODE = re.compile(r"(?:发票|票据)代码[:：\s]*([A-Z0-9\-]{8,20})")
RE_INVOICE_NO = re.compile(r"(?:发票|票据|印刷序号|票\s*号)\s*(?:号码)?[:：\s]*([0-9\-]{6,25})")
# 开票日期：2026-03-15 / 2026年03月15日 / 2026/3/15
RE_DATE = re.compile(
    r"(?:开票日期|开具日期|填开日期|日期)[:：\s]*"
    r"(\d{4})[-年./](\d{1,2})[-月./](\d{1,2})"
)

# ---- 金额正则（v1.2 全部重写，修复截断 bug）----
#
# 核心思路：金额必须是"完整的数字串"，不允许正则中的 {1,3} 造成截断。
# 采用策略：在关键锚点后，用 r"¥?\s*([\d,]+\.\d{2})" 严格匹配"整数部分+两位小数"
# 其中 [\d,]+ 是**整个整数部分**，不做位数限制（金额可能到 8 位整数）。
# 两位小数强制，避免把税率行 "13% 1225.79" 的 "13" 误当金额。

# 价税合计 · 优先级 1：（小写）¥XXX 或 小写 ¥XXX
RE_TOTAL_XIAOXIE = re.compile(
    r"(?:（小写）|\(小写\)|小写)\s*[￥¥]?\s*([\d,]+\.\d{2})"
)
# 价税合计 · 优先级 2：价税合计...¥XXX 的最后一个金额
RE_TOTAL_JIASHUI = re.compile(
    r"价税合计[\s\S]{0,120}?[￥¥]\s*([\d,]+\.\d{2})"
)
# 价税合计 · 优先级 3：飞机行程单用 "合计 TOTAL ¥XXX"
RE_TOTAL_AIR = re.compile(
    r"合计\s*TOTAL\s*[￥¥]\s*([\d,]+\.\d{2})"
)
# 价税合计 · 优先级 4：捐赠票据 "金额合计（大写）...（小写）XXX"
RE_TOTAL_DONATION = re.compile(
    r"金额合计[\s\S]{0,120}?[￥¥]?\s*([\d,]+\.\d{2})"
)
# 通用金额候选（用于兜底扫描，提取全文所有合法金额）
RE_ANY_AMOUNT = re.compile(r"[￥¥]\s*([\d,]+\.\d{2})")

# 税率：3% / 6% / 9% / 13% / 免税
RE_TAX_RATE = re.compile(r"(\d{1,2}%|免税|不征税)")
# 税额
RE_TAX_AMOUNT = re.compile(r"(?:税额|合计税额)[:：\s]*[￥¥]?\s*([\d,]+\.\d{2})")
# 购买方/销售方
RE_BUYER_NAME = re.compile(
    r"(?:购买方[\s\S]{0,5}名称|购方名称|捐赠方名称|购票单位|旅客姓名)[:：\s]*([^\n\r]+)"
)
RE_SELLER_NAME = re.compile(
    r"(?:销售方[\s\S]{0,5}名称|销方名称|受赠方名称|收款方名称|承运人)[:：\s]*([^\n\r]+)"
)
RE_TAXPAYER_ID = re.compile(r"(?:纳税人识别号|统一社会信用代码|税号)[:：\s]*([A-Z0-9]{15,20})")

# 飞机行程单专用字段
RE_AIR_ROUTE = re.compile(r"始发站\s*FROM\s*([^\n]+?)\s*目的站\s*TO\s*([^\n]+)")
RE_AIR_DATE = re.compile(r"承运日期\s*DATE\s*(\d{4})[-年./](\d{1,2})[-月./](\d{1,2})")
RE_AIR_FLIGHT = re.compile(r"航班号\s*FLIGHT\s*([A-Z0-9]+)")
RE_AIR_TICKET_NO = re.compile(r"票号\s*TICKET\s*NO\s*([\d\-]+)")


def extract_pdf_fields(pdf_path):
    """从 PDF 提取发票要素。返回 dict，失败时对应字段为空。"""
    fields = {
        "发票代码": "",
        "发票号码": "",
        "开票日期": "",
        "购买方/捐赠方": "",
        "销售方/受赠方": "",
        "税率": "",
        "税额": "",
        "价税合计": "",
        "_raw_text": "",
        "_parse_status": "未解析",
    }
    if not HAS_PDFPLUMBER:
        fields["_parse_status"] = "缺少pdfplumber"
        return fields

    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    except Exception as e:
        fields["_parse_status"] = f"解析失败: {e}"
        return fields

    # 清理 PDF 提取常见的噪声字符：空字节、非断行空格
    text = text.replace("\x00", " ").replace("\u00a0", " ").replace("\u3000", " ")

    fields["_raw_text"] = text
    fields["_parse_status"] = "成功"

    # 发票代码
    m = RE_INVOICE_CODE.search(text)
    if m:
        fields["发票代码"] = m.group(1)

    # 发票号码
    m = RE_INVOICE_NO.search(text)
    if m:
        fields["发票号码"] = m.group(1)

    # 开票日期
    m = RE_DATE.search(text)
    if m:
        y, mo, d = m.groups()
        try:
            fields["开票日期"] = f"{y}-{int(mo):02d}-{int(d):02d}"
        except ValueError:
            pass

    # 税率（取文本中最显著的一个百分比或"免税"）
    tax_rates = RE_TAX_RATE.findall(text)
    if tax_rates:
        # 优先取第一个非 100% 的税率
        for r in tax_rates:
            if r not in ("100%",):
                fields["税率"] = r
                break

    # 税额
    m = RE_TAX_AMOUNT.search(text)
    if m:
        fields["税额"] = m.group(1).replace(",", "")

    # 价税合计：4 级兜底（v1.2 重写金额正则，修复截断 bug）
    total = ""
    # 1) 飞机行程单 "合计 TOTAL ¥XXX"
    m = RE_TOTAL_AIR.search(text)
    if m:
        total = m.group(1).replace(",", "")
    # 2) "（小写）" 后面的数字
    if not total:
        m = RE_TOTAL_XIAOXIE.search(text)
        if m:
            total = m.group(1).replace(",", "")
    # 3) "价税合计...¥XXX"
    if not total:
        m = RE_TOTAL_JIASHUI.search(text)
        if m:
            total = m.group(1).replace(",", "")
    # 4) 捐赠票据 "金额合计...¥XXX"
    if not total:
        m = RE_TOTAL_DONATION.search(text)
        if m:
            total = m.group(1).replace(",", "")
    # 5) 最终兜底：取全文中最后一个 ¥XXX（通常是价税合计）
    if not total:
        all_amounts = RE_ANY_AMOUNT.findall(text)
        if all_amounts:
            # 取最后一个（价税合计通常在文本靠后位置）
            total = all_amounts[-1].replace(",", "")

    if total:
        fields["价税合计"] = total

    # 购方/销方
    m = RE_BUYER_NAME.search(text)
    if m:
        fields["购买方/捐赠方"] = m.group(1).strip()
    m = RE_SELLER_NAME.search(text)
    if m:
        fields["销售方/受赠方"] = m.group(1).strip()

    return fields


# ============================================================
# DOCX 文本提取（v1.2 新增）
# ============================================================

def extract_docx_fields(docx_path):
    """从 DOCX（打印版发票/报销单模板）提取文本和字段。"""
    fields = {
        "发票代码": "",
        "发票号码": "",
        "开票日期": "",
        "购买方/捐赠方": "",
        "销售方/受赠方": "",
        "税率": "",
        "税额": "",
        "价税合计": "",
        "_raw_text": "",
        "_parse_status": "未解析",
    }
    if not HAS_DOCX:
        fields["_parse_status"] = "缺少python-docx"
        return fields
    try:
        doc = docx.Document(str(docx_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也读表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        fields["_raw_text"] = text
        fields["_parse_status"] = "成功"

        # 复用 PDF 的正则（字段名一样）
        m = RE_INVOICE_CODE.search(text)
        if m:
            fields["发票代码"] = m.group(1)
        m = RE_INVOICE_NO.search(text)
        if m:
            fields["发票号码"] = m.group(1)
        m = RE_DATE.search(text)
        if m:
            y, mo, d = m.groups()
            try:
                fields["开票日期"] = f"{y}-{int(mo):02d}-{int(d):02d}"
            except ValueError:
                pass
        m = RE_TAX_RATE.search(text)
        if m:
            fields["税率"] = m.group(1)
        m = RE_TAX_AMOUNT.search(text)
        if m:
            fields["税额"] = m.group(1).replace(",", "")
        # 金额 4 级兜底
        total = ""
        for pat in (RE_TOTAL_AIR, RE_TOTAL_XIAOXIE, RE_TOTAL_JIASHUI, RE_TOTAL_DONATION):
            if not total:
                m = pat.search(text)
                if m:
                    total = m.group(1).replace(",", "")
        if not total:
            all_amt = RE_ANY_AMOUNT.findall(text)
            if all_amt:
                total = all_amt[-1].replace(",", "")
        if total:
            fields["价税合计"] = total
        m = RE_BUYER_NAME.search(text)
        if m:
            fields["购买方/捐赠方"] = m.group(1).strip()
        m = RE_SELLER_NAME.search(text)
        if m:
            fields["销售方/受赠方"] = m.group(1).strip()
    except Exception as e:
        fields["_parse_status"] = f"解析失败: {e}"
    return fields


# ============================================================
# 图片 OCR 提取（v3 新增）
# ============================================================

# 图片票据通用信息正则（OCR 文本可能有错字，宽松匹配）
RE_OCR_DATE = re.compile(
    r"(20\d{2})[\-\./年](\d{1,2})[\-\./月](\d{1,2})"
)
RE_OCR_AMOUNT = re.compile(
    # 匹配 ¥/口/0 等误识字符 + 金额
    r"(?:[￥¥口0O]|金额|合计|应付|票价|总计)\s*[:：]?\s*"
    r"([0-9]{1,3}(?:,\d{3})*(?:\.\d{1,2})?|\d+\.\d{1,2})"
)
# 纯金额候选（作为兜底，在全文里找所有带小数点的数字）
RE_OCR_NUM = re.compile(r"(\d{1,3}(?:,\d{3})+(?:\.\d{1,2})?|\d+\.\d{2})")


def _load_image_as_array(img_path):
    """用 PIL 加载图片并转 numpy 数组，规避 OpenCV 不支持中文路径的问题"""
    try:
        from PIL import Image
        import numpy as np
        return np.array(Image.open(str(img_path)).convert("RGB"))
    except Exception:
        return None


def extract_image_fields(img_path):
    """从图片票据 OCR 提取关键信息。返回字段结构与 extract_pdf_fields 一致。"""
    fields = {
        "发票代码": "",
        "发票号码": "",
        "开票日期": "",
        "购买方/捐赠方": "",
        "销售方/受赠方": "",
        "税率": "",
        "税额": "",
        "价税合计": "",
        "_raw_text": "",
        "_parse_status": "未解析",
    }

    reader = _get_ocr_reader()
    if reader is None:
        fields["_parse_status"] = "缺少easyocr"
        return fields

    arr = _load_image_as_array(img_path)
    if arr is None:
        fields["_parse_status"] = "图片读取失败"
        return fields

    try:
        lines = reader.readtext(arr, detail=0, paragraph=False)
    except Exception as e:
        fields["_parse_status"] = f"OCR失败: {e}"
        return fields

    text = "\n".join(lines)
    fields["_raw_text"] = text
    fields["_parse_status"] = "OCR成功"

    # 日期
    m = RE_OCR_DATE.search(text)
    if m:
        y, mo, d = m.groups()
        try:
            fields["开票日期"] = f"{y}-{int(mo):02d}-{int(d):02d}"
        except ValueError:
            pass

    # 价税合计策略（v3.1 改进）：
    # 1) 优先在"合计/应付/票价/金额/总计"关键词**同行或下一行**找金额
    # 2) 排除过长的数字串（>10 位大概率是票号/单号，不是金额）
    # 3) 若关键词附近无匹配，退回取次大的合理金额（排除票号干扰）
    KEYWORDS_TOTAL = ("合计", "应付", "票价", "票面", "总计", "合 计", "合價", "金額合計")
    lines_split = [ln.strip() for ln in text.split("\n") if ln.strip()]
    total_amount = None

    def _parse_num(s):
        try:
            v = float(s.replace(",", "").replace("，", ""))
            # 金额合理范围：0.01 ~ 9_999_999.99
            if 0.01 <= v <= 9_999_999.99:
                return v
        except ValueError:
            pass
        return None

    def _looks_like_date(s, context_text):
        """判断候选字符串是否像日期的一部分（如 '2026.03' 其实是 2026.03.12 的前缀）"""
        if not s:
            return False
        # 单独的年份（2020~2099）
        try:
            if s.replace(".", "").replace(",", "").isdigit():
                v = float(s.replace(",", ""))
                if 2020 <= v <= 2099 and "." not in s:
                    return True
        except ValueError:
            pass
        # 字符串形如 "2026.03" / "2026.3" 且原文中后面紧跟 ".dd"
        idx = context_text.find(s)
        if idx >= 0:
            tail = context_text[idx + len(s): idx + len(s) + 4]
            if re.match(r"\.\d{1,2}", tail):
                return True
        return False

    # 第 1 步：在关键词行内/下一行查找
    for i, ln in enumerate(lines_split):
        if any(kw in ln for kw in KEYWORDS_TOTAL):
            search_text = ln + " " + (lines_split[i + 1] if i + 1 < len(lines_split) else "")
            # 优先匹配带小数点的金额
            cands = re.findall(r"\d{1,6}\.\d{2}", search_text)
            if not cands:
                cands = re.findall(r"\d{1,3}(?:,\d{3})+(?:\.\d+)?", search_text)
            if not cands:
                cands = re.findall(r"\b\d{2,6}\b", search_text)
            for c in cands:
                if _looks_like_date(c, text):
                    continue
                v = _parse_num(c)
                if v:
                    total_amount = v
                    break
            if total_amount:
                break

    # 第 2 步：若关键词未命中，找全文所有"合理金额"（带小数点/千分号），取最大
    if not total_amount:
        candidates = []
        for m in re.finditer(r"\d{1,3}(?:,\d{3})+\.\d{2}|\d{1,3}(?:,\d{3})+|\d+\.\d{2}", text):
            s = m.group(0)
            if _looks_like_date(s, text):
                continue
            v = _parse_num(s)
            if v:
                candidates.append(v)
        if candidates:
            total_amount = max(candidates)

    if total_amount:
        fields["价税合计"] = f"{total_amount:.2f}"

    # 票号：找 8~20 位连续数字（排除日期和金额）
    date_digits = {fields["开票日期"].replace("-", "")} if fields["开票日期"] else set()
    for token in re.findall(r"\b(\d{8,20})\b", text):
        if token in date_digits:
            continue
        # 排除像 20250225 这种可能是日期的数字
        if len(token) == 8 and token.startswith("20") and 1 <= int(token[4:6]) <= 12:
            continue
        fields["发票号码"] = token
        break

    return fields


# ============================================================
# 分类与文件名元数据提取
# ============================================================

def classify_file(file_path, file_text=""):
    """
    根据文件名、文件内容关键词分类。
    file_text 是已解析的文本（PDF/DOCX/图片 OCR 结果），优先用内容分类。
    返回"细分类别"，由调用方映射到归档目录。
    """
    name_lower = file_path.stem.lower() + " " + str(file_path.parent).lower()
    # 同时搜文件名 + 文件内容（如果提供了）
    haystack = (name_lower + " " + file_text.lower()) if file_text else name_lower

    # 按关键词长度倒序匹配（特异性优先）
    for keyword, category in sorted(CATEGORY_KEYWORDS, key=lambda x: len(x[0]), reverse=True):
        if keyword.lower() in haystack:
            return category

    ext = file_path.suffix.lower()
    if ext in EXCEL_EXTS:
        return "Excel数据表"
    if ext in DOCX_EXTS:
        return "DOCX文件"
    if ext in PDF_EXTS:
        return "PDF文件"
    if ext in IMAGE_EXTS:
        return "未分类票据图片"
    return "其他"


def suggest_expense_category(file_path, classification=""):
    """根据文件名和分类猜测费用科目"""
    name = file_path.stem.lower()
    if "捐赠" in classification:
        return "捐赠收入—限定性/非限定性（需判断）"
    for keyword, category in EXPENSE_KEYWORDS.items():
        if keyword in name:
            return category
    return "待分类"


def extract_date_from_name(file_path):
    """从文件名提取日期"""
    name = file_path.stem
    patterns = [
        r"(\d{4})[-_.](\d{1,2})[-_.](\d{1,2})",
        r"(\d{4})(\d{2})(\d{2})",
    ]
    for pattern in patterns:
        match = re.search(pattern, name)
        if match:
            groups = match.groups()
            try:
                return f"{groups[0]}-{int(groups[1]):02d}-{int(groups[2]):02d}"
            except ValueError:
                pass
    return ""


def extract_amount_from_name(file_path):
    """从文件名提取金额"""
    name = file_path.stem
    patterns = [
        r"(\d+(?:\.\d{1,2})?)\s*元",
        r"¥\s*(\d+(?:\.\d{1,2})?)",
    ]
    for pattern in patterns:
        match = re.search(pattern, name, re.IGNORECASE)
        if match:
            return float(match.group(1))
    return None


# ============================================================
# 文件整理
# ============================================================

def organize_files(files, output_folder):
    """将文件按分类复制到输出文件夹，同时提取发票要素"""
    output = Path(output_folder)
    output.mkdir(parents=True, exist_ok=True)
    organized = []

    for f in files:
        name_date = extract_date_from_name(f)
        name_amount = extract_amount_from_name(f)

        # 按格式提取字段
        file_fields = {}
        raw_text = ""
        ext = f.suffix.lower()
        if ext in PDF_EXTS:
            file_fields = extract_pdf_fields(f)
        elif ext in IMAGE_EXTS:
            file_fields = extract_image_fields(f)
        elif ext in DOCX_EXTS:
            file_fields = extract_docx_fields(f)
        raw_text = file_fields.get("_raw_text", "")
        parse_status = file_fields.get("_parse_status", "未解析")

        # 用文件名 + 内容做分类（内容优先）
        category = classify_file(f, raw_text)
        expense_cat = suggest_expense_category(f, category)

        # 整合字段（抽取到的优先，文件名做兜底）
        invoice_code = file_fields.get("发票代码", "") or ""
        invoice_no = file_fields.get("发票号码", "") or ""
        issue_date = file_fields.get("开票日期", "") or name_date or ""
        buyer = file_fields.get("购买方/捐赠方", "") or ""
        seller = file_fields.get("销售方/受赠方", "") or ""
        tax_rate = file_fields.get("税率", "") or ""
        tax_amt = file_fields.get("税额", "") or ""
        total_amt = file_fields.get("价税合计", "") or (str(name_amount) if name_amount else "")

        # 数据状态判断（诚实原则）
        # 图片/OCR 结果 → 默认需人工复核
        # PDF/DOCX 缺核心字段（金额、日期、号码任一缺失）→ 也需人工复核
        core_missing = not (total_amt and issue_date and invoice_no)
        if ext in IMAGE_EXTS:
            data_status = "需人工复核（OCR）"
        elif core_missing:
            data_status = "需人工复核（缺字段）"
        else:
            data_status = "待复核"

        # 分类归档到标准目录
        archive_dir = CATEGORY_ARCHIVE_DIR.get(category, "06_其他")
        cat_folder = output / archive_dir
        cat_folder.mkdir(parents=True, exist_ok=True)
        new_name = f"{issue_date}_{f.name}" if issue_date else f.name
        dest = cat_folder / new_name
        counter = 1
        while dest.exists():
            dest = cat_folder / f"{dest.stem}_{counter}{dest.suffix}"
            counter += 1
        shutil.copy2(str(f), str(dest))

        organized.append({
            "序号": len(organized) + 1,
            "票据类型": category,
            "票据名称": f.stem,
            "发票代码": invoice_code,
            "发票号码": invoice_no,
            "开票日期": issue_date,
            "购买方/捐赠方": buyer,
            "销售方/受赠方": seller,
            "货物或服务项目": "",  # 项目明细一般较长，需人工补录
            "不含税金额": "",
            "税率/税额": f"{tax_rate}/{tax_amt}" if tax_rate or tax_amt else "",
            "价税合计(元)": total_amt,
            "建议科目": expense_cat,
            "备注": "自动提取" if data_status == "待复核" else f"[{parse_status}] {data_status}",
            "数据状态": data_status,
            "原文件路径": str(f),
            "新文件路径": str(dest),
            "文件类型": f.suffix.lower(),
        })

    return organized


# ============================================================
# 台账输出（Excel 优先，CSV 兜底）
# ============================================================

LEDGER_COLS = [
    "序号", "票据类型", "票据名称",
    "发票代码", "发票号码", "开票日期",
    "购买方/捐赠方", "销售方/受赠方",
    "货物或服务项目", "不含税金额", "税率/税额", "价税合计(元)",
    "建议科目", "备注", "数据状态",
]


def generate_ledger_xlsx(organized, output_folder):
    """生成 Excel 台账（15 列，含完整发票要素）"""
    if not HAS_OPENPYXL:
        return None
    output = Path(output_folder)
    ledger_path = output / "票据台账_自动生成.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "发票台账"

    # 样式
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="1F4E78")
    border = Border(
        left=Side(style="thin", color="BFBFBF"),
        right=Side(style="thin", color="BFBFBF"),
        top=Side(style="thin", color="BFBFBF"),
        bottom=Side(style="thin", color="BFBFBF"),
    )
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # 标题行
    ws.cell(row=1, column=1, value=f"发票台账（自动生成 {datetime.now():%Y-%m-%d %H:%M}）")
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(LEDGER_COLS))
    ws.cell(row=1, column=1).font = Font(bold=True, size=13)
    ws.cell(row=1, column=1).alignment = center

    # 表头
    for i, col in enumerate(LEDGER_COLS, 1):
        c = ws.cell(row=2, column=i, value=col)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center
        c.border = border

    # 数据
    for row_idx, item in enumerate(organized, 3):
        for col_idx, col in enumerate(LEDGER_COLS, 1):
            value = item.get(col, "")
            # 数值列转 float
            if col == "价税合计(元)" and value:
                try:
                    value = float(value)
                except (ValueError, TypeError):
                    pass
            c = ws.cell(row=row_idx, column=col_idx, value=value)
            c.alignment = left_wrap if col in ("购买方/捐赠方", "销售方/受赠方",
                                                "货物或服务项目", "备注") else center
            c.border = border
            # 需人工复核的行标黄
            if "需人工复核" in item.get("数据状态", ""):
                c.fill = PatternFill("solid", fgColor="FFF2CC")

    # 合计行
    total_row = len(organized) + 3
    ws.cell(row=total_row, column=1, value="合计")
    ws.cell(row=total_row, column=1).font = Font(bold=True)
    total_sum = sum(
        float(x["价税合计(元)"]) for x in organized
        if x.get("价税合计(元)") and str(x["价税合计(元)"]).replace(".", "").replace("-", "").isdigit()
    )
    ws.cell(row=total_row, column=12, value=total_sum).font = Font(bold=True)
    for i in range(1, len(LEDGER_COLS) + 1):
        ws.cell(row=total_row, column=i).fill = PatternFill("solid", fgColor="D9E1F2")
        ws.cell(row=total_row, column=i).border = border

    # 列宽
    widths = [6, 18, 22, 16, 14, 12, 28, 28, 32, 12, 12, 14, 24, 22, 12]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.row_dimensions[2].height = 30

    # 冻结首行
    ws.freeze_panes = "A3"

    wb.save(str(ledger_path))
    return ledger_path


def generate_ledger_csv(organized, output_folder):
    """CSV 兜底台账"""
    output = Path(output_folder)
    ledger_path = output / "票据台账_自动生成.csv"
    with open(str(ledger_path), "w", newline="", encoding="utf-8-sig") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=LEDGER_COLS)
        writer.writeheader()
        for item in organized:
            writer.writerow({k: item.get(k, "") for k in LEDGER_COLS})
    return ledger_path


# ============================================================
# 汇总报告
# ============================================================

def generate_summary(organized, output_folder, ledger_path):
    output = Path(output_folder)
    summary_path = output / "整理汇总报告.md"

    by_category = defaultdict(list)
    for item in organized:
        by_category[item["票据类型"]].append(item)

    total_files = len(organized)
    pdf_parsed_ok = sum(1 for x in organized if "自动提取" in x.get("备注", ""))
    need_manual = sum(1 for x in organized if "需人工复核" in x.get("数据状态", ""))
    total_amt = sum(
        float(x["价税合计(元)"]) for x in organized
        if x.get("价税合计(元)") and str(x["价税合计(元)"]).replace(".", "").replace("-", "").isdigit()
    )

    lines = [
        "# 票据自动整理汇总报告",
        "",
        "> ⚠️ **警示**：本报告及配套台账中的金额、日期、发票号均由脚本自动提取，使用前必须人工逐张复核，复核通过后方可入账。",
        "",
        f"> 生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}",
        f"> 台账文件：`{Path(ledger_path).name}`" if ledger_path else "",
        "",
        "---",
        "",
        "## 一、总体概况",
        "",
        "| 指标 | 数值 |",
        "|------|------|",
        f"| 扫描文件总数 | **{total_files}** 张 |",
        f"| 字段自动提取成功 | {pdf_parsed_ok} 张 |",
        f"| 需人工复核（OCR / 缺字段） | {need_manual} 张 |",
        f"| 已抓取金额合计 | **¥{total_amt:,.2f}** |",
        "",
        "## 二、分类明细",
        "",
        "| 票据类型 | 张数 |",
        "|---------|:---:|",
    ]
    for cat, items in sorted(by_category.items()):
        lines.append(f"| {cat} | {len(items)} |")

    lines.extend([
        "",
        "## 三、需人工处理事项",
        "",
    ])

    manual_items = [x for x in organized if "需人工复核" in x.get("数据状态", "")]
    if manual_items:
        lines.append(f"### ⚠️ 以下 {len(manual_items)} 张票据需人工复核（OCR 或缺字段）：")
        lines.append("")
        lines.append("| 序号 | 文件名 | 票据类型 | 数据状态 | 建议科目 |")
        lines.append("|:-:|---|---|---|---|")
        for x in manual_items:
            lines.append(f"| {x['序号']} | {x['票据名称']} | {x['票据类型']} | {x['数据状态']} | {x['建议科目']} |")
        lines.append("")
        lines.append("> 复核要点：")
        lines.append("> 1. 金额：对照原票小数点、千分位、正负号")
        lines.append("> 2. 日期：注意年份跨度和格式（YYYY-MM-DD）")
        lines.append("> 3. 发票号：18 位纳税人识别号、8~20 位发票号码")
        lines.append("> 4. 复核后在台账「数据状态」列改为「已复核」")
        lines.append("")
    else:
        lines.append("✅ 本次整理所有票据核心字段均已自动提取，但仍建议逐行核对金额和日期。")
        lines.append("")

    lines.extend([
        "## 四、下一步建议",
        "",
        "1. 打开 Excel 台账，逐行复核「发票代码、发票号码、开票日期」三个关键字段",
        "2. 将「数据状态」=「需人工复核」的行优先处理",
        "3. 增值税发票上国税平台查验（https://inv-veri.chinatax.gov.cn）",
        "4. 捐赠票据登记到《票据使用情况汇总表》，年末向财政部门申报核销",
        "",
        "---",
        "",
        "> 💡 本报告由 公益票据自动整理脚本 v1.2 生成。PDF 解析基于 pdfplumber，图片 OCR 基于 easyocr，DOCX 基于 python-docx。如需更高精度，建议使用腾讯云 OCR API。",
    ])

    with open(str(summary_path), "w", encoding="utf-8") as f:
        f.write("\n".join(x for x in lines if x is not None))

    return summary_path


# ============================================================
# 主流程
# ============================================================

def main():
    # Windows PowerShell 默认 GBK 不支持部分 emoji，强制 UTF-8
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

    import argparse
    parser = argparse.ArgumentParser(
        description="票据自动扫描与分类整理工具 v1.2（含发票要素自动提取）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("folder", help="要扫描的票据文件夹路径")
    parser.add_argument("--output", "-o", help="输出文件夹路径")
    parser.add_argument("--scan-only", action="store_true", help="仅扫描统计，不复制文件")
    args = parser.parse_args()

    folder_path = args.folder
    if not args.output:
        output_folder = os.path.join(
            folder_path,
            f"整理结果_{datetime.now():%Y%m%d_%H%M%S}"
        )
    else:
        output_folder = args.output

    print("=" * 60)
    print("[INFO] 票据自动扫描与分类整理工具 v1.2")
    print("=" * 60)
    print(f"扫描目录: {folder_path}")
    if not HAS_PDFPLUMBER:
        print("  [WARN] pdfplumber 未安装，PDF 字段提取将被跳过。建议: pip install pdfplumber")
    if not HAS_OPENPYXL:
        print("  [WARN] openpyxl 未安装，将回退到 CSV 输出。建议: pip install openpyxl")
    if not HAS_DOCX:
        print("  [WARN] python-docx 未安装，DOCX 支持将被跳过。建议: pip install python-docx")
    print()

    # 扫描
    files = []
    folder = Path(folder_path)
    if not folder.exists():
        print(f"❌ 文件夹不存在: {folder_path}")
        sys.exit(1)
    # 排除模式：脚本生成物、台账类 Excel、临时文件
    EXCLUDE_PATTERNS = [
        "整理结果_", "票据台账", "台账_", "统计表", "汇总报告",
        "~$",  # Office 临时文件
    ]
    for f in folder.rglob("*"):
        skip = any(p in str(f) for p in EXCLUDE_PATTERNS)
        if skip:
            continue
        if f.is_file() and f.suffix.lower() in ALL_SUPPORTED:
            files.append(f)
    print(f"🔍 找到 {len(files)} 个支持的文件")

    # 分类预览
    preview = defaultdict(int)
    for f in files:
        preview[classify_file(f)] += 1
    print()
    print("📊 分类预览:")
    for cat, count in sorted(preview.items()):
        print(f"   {cat}: {count} 个")

    if args.scan_only:
        print()
        print("✅ 扫描完成（仅扫描模式）")
        return

    # 整理 + 提取
    print()
    print(f"📂 正在整理到: {output_folder}")
    organized = organize_files(files, output_folder)
    print(f"   归档 {len(organized)} 个文件")

    # 生成台账
    print()
    print("📝 正在生成台账...")
    if HAS_OPENPYXL:
        ledger_path = generate_ledger_xlsx(organized, output_folder)
    else:
        ledger_path = generate_ledger_csv(organized, output_folder)
    print(f"   台账: {ledger_path}")

    # 汇总报告
    summary_path = generate_summary(organized, output_folder, ledger_path)
    print(f"   报告: {summary_path}")

    # 解析统计
    pdf_ok = sum(1 for x in organized if "自动提取" in x.get("备注", ""))
    need_manual = sum(1 for x in organized if "需人工复核" in x.get("数据状态", ""))
    print()
    print("=" * 60)
    print("✅ 整理完成！")
    print(f"   📂 结果目录: {output_folder}")
    print(f"   📊 字段自动提取: {pdf_ok} 张")
    print(f"   ⚠️  需人工复核: {need_manual} 张（OCR / 缺字段）")
    print("   ⚠️  脚本金额必须人工逐张复核后方可入账")
    print("=" * 60)


if __name__ == "__main__":
    main()
