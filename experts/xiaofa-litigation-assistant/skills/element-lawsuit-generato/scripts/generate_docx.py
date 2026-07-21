#!/usr/bin/env python3
"""
要素式起诉状 DOCX 生成器
数据驱动，基于 法〔2025〕82号 官方模板 table_layouts.json。
精确复刻所有33个民事起诉状案由的表格结构。
"""
import sys, os, json, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══ Paths ═══
SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LAYOUTS_PATH = os.path.join(SKILL_DIR, 'references', 'table_layouts.json')
TEMPLATES_PATH = os.path.join(SKILL_DIR, 'references', 'templates.json')

# ═══ Constants (from official 82号 template) ═══
TABLE_W = 9344
COL_L   = 2270
COL_R   = 7074
FONT    = '宋体'
SZ      = Pt(10.5)
SZ_HDR  = Pt(15)
SZ_TTL  = Pt(22)
SZ_SUB  = Pt(18)

# Page setup (A4, official margins from 法〔2025〕82号)
PAGE_W = 11906
PAGE_H = 16838
MARGIN_TOP    = 908685
MARGIN_BOTTOM = 633730
MARGIN_LEFT   = 899795
MARGIN_RIGHT  = 719455

# ═══ Load layouts ═══
_table_layouts = {}
_templates = {}

def _load_data():
    global _table_layouts, _templates
    if not _table_layouts:
        if os.path.exists(LAYOUTS_PATH):
            with open(LAYOUTS_PATH, 'r', encoding='utf-8') as f:
                _table_layouts = json.load(f)
    if not _templates:
        if os.path.exists(TEMPLATES_PATH):
            with open(TEMPLATES_PATH, 'r', encoding='utf-8') as f:
                _templates = json.load(f)


# ═══ XML helpers ═══

def _oxml(tag, attrs=None):
    if attrs is None:
        attrs = {}
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = [f'<w:{tag}']
    for k, v in attrs.items():
        parts.append(f' w:{k}="{v}"')
    parts.append('/>')
    return ''.join(parts)

def _make_borders():
    b = _oxml
    sides = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    inner = ''.join(b(s, {'val': 'single', 'color': '231F20', 'sz': '2', 'space': '0'}) for s in sides)
    return f'<w:tblBorders {nsdecls("w")}>{inner}</w:tblBorders>'

def _new_table(doc):
    table = doc.add_table(rows=0, cols=2)
    tbl = table._tbl
    tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{TABLE_W}" w:type="dxa"/>'))
    tblPr.append(parse_xml(_make_borders()))
    tbl.insert(0, tblPr)
    return table

def _run(para, text, bold=False, size=SZ, font=FONT):
    r = para.add_run(str(text))
    r.font.size = size
    r.font.name = font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    r.bold = bold
    return r

def _cell_setup(cell, width):
    for p in cell.paragraphs:
        p.clear()
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width}" w:type="dxa"/>'))
    for old in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))

def _merged_row(table, text, size=SZ):
    row = table.add_row()
    cell = row.cells[0]
    cell.merge(row.cells[1])
    _cell_setup(cell, TABLE_W)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=False, size=size)

def _label_data_row(table, label, data_lines, lbl_w=None, dat_w=None):
    if lbl_w is None: lbl_w = COL_L
    if dat_w is None: dat_w = COL_R
    row = table.add_row()
    lc = row.cells[0]
    _cell_setup(lc, lbl_w)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(lp, label)
    dc = row.cells[1]
    _cell_setup(dc, dat_w)
    for i, line in enumerate(data_lines):
        dp = dc.paragraphs[0] if i == 0 else dc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(dp, str(line))

def _add_instruction_merged(table, text):
    """Add the 说明 instruction as separate paragraphs within a merged cell."""
    row = table.add_row()
    cell = row.cells[0]
    cell.merge(row.cells[1])
    _cell_setup(cell, TABLE_W)
    # Split the long text into proper instruction paragraphs
    lines = [
        '说明：',
        '为了方便您更好地参加诉讼，保护您的合法权利，请填写本表。',
        '1.起诉时需向人民法院提交证明您身份的材料，如身份证复印件、营业执照复印件等。',
        '2.本表所列内容是您提起诉讼以及人民法院查明案件事实所需，请务必如实填写。',
        '3.本表所涉内容系针对一般案件，有些内容可能与您的案件无关，您认为与案件无关的项目可以填"无"或不填；对于本表中勾选项可以在对应项打"√"；您认为另有重要内容需要列明的，可以在本表尾部或者另附页填写。',
        '★特别提示★',
        '《中华人民共和国民事诉讼法》第十三条第一款规定："民事诉讼应当遵循诚信原则。"',
        '如果诉讼参加人违反上述规定，进行虚假诉讼、恶意诉讼，人民法院将视违法情形依法追究责任。',
    ]
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(para, line, bold=(i == 0))

# ═══ Content formatters ═══

def _fmt_org(p, prefix=''):
    lines = []
    lines.append(f'{prefix}名称：{p.get("name","") or ""}')
    lines.append(f'住所地（主要办事机构所在地）：{p.get("addr","") or ""}')
    lines.append(f'注册地/登记地：{p.get("regAddr","") or ""}')
    lr = p.get('legalRep', '') or ''
    job = p.get('job', '') or ''
    phone = p.get('phone', '') or ''
    lines.append(f'法定代表人/主要负责人：{lr}  职务：{job}  联系电话：{phone}')
    lines.append(f'统一社会信用代码：{p.get("creditCode","") or ""}')
    types = ['有限责任公司','股份有限公司','上市公司','其他企业法人',
             '事业单位','社会团体','基金会','社会服务机构','机关法人',
             '农村集体经济组织法人','城镇农村的合作经济组织法人',
             '基层群众性自治组织法人','个人独资企业','合伙企业',
             '不具有法人资格的专业服务机构']
    ot = p.get('orgType', '') or ''
    checked = [f'{"☑" if t == ot else "□"} {t}' for t in types]
    for g in range(0, len(checked), 4):
        lines.append('  '.join(checked[g:g+4]))
    owner = p.get('ownership', '') or ''
    if '国有' in owner:
        gk = '☑' if '控股' in owner else '□'
        cg = '☑' if '参股' in owner else '□'
        lines.append(f'国有☑ （控股{gk}参股{cg}）民营□ 其他□')
    elif '民营' in owner:
        lines.append('国有□ （控股□参股□）民营☑ 其他□')
    else:
        lines.append('国有□ （控股□参股□）民营□ 其他☑')
    return lines

def _fmt_person(p, prefix=''):
    g = p.get('gender', '') or ''
    idn = p.get('idNum', '') or ''
    lines = [f'{prefix}姓名：{p.get("name","") or ""}']
    if g == '男':
        gs = '性别：男☑ 女□'
    elif g == '女':
        gs = '性别：男□ 女☑'
    else:
        gs = '性别：男□ 女□'
    lines.append(f'{gs}  身份证号码：{idn}')
    birth = (p.get('birth', '') or '')
    nation = p.get('nation', '') or ''
    lines.append(f'出生日期：{birth}        民族：{nation}')
    work = p.get('work', '') or ''
    job = p.get('job', '') or ''
    phone = p.get('phone', '') or ''
    lines.append(f'工作单位：{work}  职务：{job}  联系电话：{phone}')
    lines.append(f'住所地（户籍所在地）：{p.get("addr","") or ""}')
    habitual = p.get('habitual', '') or ''
    if habitual:
        lines.append(f'送达地址：{habitual}')
    return lines

def _fmt_agent(a):
    has = a.get('has', False)
    chk = '☑' if has else '□'
    nochk = '□' if has else '☑'
    return [
        f'有{chk}    无{nochk}',
        f'姓名：{a.get("name","") or ""}    单位：{a.get("firm","") or ""}    职务：{a.get("job","") or ""}',
        f'联系电话：{a.get("phone","") or ""}',
        f'代理权限：一般授权{"☑" if a.get("auth","general") == "general" else "□"}    特别授权{"☑" if a.get("auth") == "special" else "□"}',
    ]

def _empty_person():
    return [
        '姓名：/', '性别：男□ 女□',
        '出生日期：   年  月   日        民族：/',
        '工作单位：/      职务：/      联系电话：/',
        '住所地（户籍所在地）：/',
    ]

def _empty_org():
    return [
        '名称：/', '住所地（主要办事机构所在地）：/',
        '注册地/登记地：/', '法定代表人/主要负责人：/  职务：/  联系电话：/',
        '统一社会信用代码：/',
    ]

def _empty_agent():
    return [
        '有□    无☑',
        '姓名：    单位：                    职务：                    联系电话：',
        '代理权限：一般授权□    特别授权□',
    ]

# ═══ Label → Data matching ═══

def _match_label_to_data(label, case_data):
    """Given a label text, return the appropriate data lines to fill in."""
    claims = case_data.get('claims', {})
    facts = case_data.get('facts', {})
    juris = case_data.get('jurisdiction', {})
    
    # Try to match as claim (numbered items)
    claim_match = re.match(r'^(\d+)\.\s*(.+)', label)
    if claim_match:
        num = int(claim_match.group(1))
        # Look up in claims data by various key patterns
        key = f'claim_{num:02d}'
        if key in claims:
            val = claims[key]
            return val.split('\n') if val else ['']
        # Also try facts
        key = f'facts_{num-1:02d}'
        if key in facts:
            val = facts[key]
            return val.split('\n') if val else ['']
    
    # Jurisdiction labels
    if '有无仲裁' in label or '法院管辖' in label:
        basis = juris.get('basis', '')
        if basis:
            return ['有☑   合同条款及内容：', basis, '无□']
        return ['有□   合同条款及内容：', '无☑']
    
    if '是否已经诉前保全' in label or '诉前保全' in label:
        pres = claims.get('preservation', juris.get('preservation', 'no'))
        if pres == 'yes':
            return ['是☑   保全法院：            保全时间：            保全案号：', '否□', '（如申请诉讼保全，请另行提交诉讼保全申请及相关材料）']
        return ['是□   保全法院：            保全时间：            保全案号：', '否☑', '（如申请诉讼保全，请另行提交诉讼保全申请及相关材料）']
    
    # Mediation labels
    if '是否了解调解作为非诉讼' in label:
        return ['了解□   不了解□']
    
    if '是否了解先行调解解决纠纷的好处' in label:
        # Return benefits 1-4 for first occurrence, benefit 5 for second
        # We'll distinguish by context in the calling code
        return ['1.立案后选择先行调解的，可以很快启动调解程序。如不同意调解，法院将依程序开庭审理案件，但可能需要经过较长一段时间才能拿到判决结果。',
                '2.先行调解不收取任何费用，如果后续达成调解协议，亦无需交纳案件受理费。',
                '3.先行调解可以选择您认为方便的方式，包括线上调解或线下调解。',
                '4.与开庭审理相比，先行调解的程序更加灵活、便捷。']
    
    if '是否考虑先行调解' in label:
        med = juris.get('mediation', 'no')
        if med == 'yes':
            return ['是☑', '否□', '暂不确定，想要了解更多内容□']
        elif med == 'unsure':
            return ['是□', '否□', '暂不确定，想要了解更多内容☑']
        return ['是□', '否☑', '暂不确定，想要了解更多内容□']
    
    # Other labels with specific formatting
    if '第三人（法人、非法人组织）' in label and '类型' not in label:
        # This is a repeat of the third party label with org type detail
        org_t = [t for t in case_data.get('thirds', []) if t.get('type') == 'org']
        if org_t:
            return [
                '类型：有限责任公司□   股份有限公司□   上市公司□',
                '其他企业法人□   事业单位□   社会团体□   基金会□   社会服务机构□   机关法人□',
                '农村集体经济组织法人□   城镇农村的合作经济组织法人□   基层群众性自治组织法人□',
                '个人独资企业□   合伙企业□   不具有法人资格的专业服务机构□',
                '国有□ （控股□参股□）民营□ 其他□',
            ]
        return ['']
    
    # Default: return empty
    return ['']

# ═══ Layout-driven generation ═══

def _render_table(doc, layout_table, case_data, data_context):
    """Render a single table from layout data, filling in case data."""
    table = _new_table(doc)
    
    # Track mediation benefit occurrence (benefits 1-4 vs benefit 5)
    mediation_benefit_count = 0
    
    for row in layout_table['rows']:
        rtype = row['type']
        
        if rtype == 'merged':
            # Check if this is the instruction row
            if '说明' in row['text'][:10]:
                _add_instruction_merged(table, row['text'])
            else:
                _merged_row(table, row['text'])
        elif rtype == 'section_header':
            _merged_row(table, row['text'], size=SZ_HDR)
        elif rtype == 'label_data':
            label = row['label']
            label_clean = re.sub(r'\s+', '', label)  # for matching
            
            # ── Fixed party rows ──
            if '原告（自然人）' in label:
                per_p = data_context['per_plaintiffs']
                if per_p:
                    _label_data_row(table, label, _fmt_person(per_p[0]))
                else:
                    _label_data_row(table, label, _empty_person())
            elif '原告（法人' in label:
                org_p = data_context['org_plaintiffs']
                if org_p:
                    _label_data_row(table, label, _fmt_org(org_p[0]))
                else:
                    _label_data_row(table, label, _empty_org())
            elif '原告' in label:
                # Generic 原告 row (离婚纠纷 uses this)
                per_p = data_context['per_plaintiffs']
                if per_p:
                    _label_data_row(table, label, _fmt_person(per_p[0]))
                else:
                    _label_data_row(table, label, _empty_person())
            elif '被告（自然人）' in label:
                per_d = data_context['per_defendants']
                org_d = data_context['org_defendants']
                if per_d:
                    nums = ['一','二','三','四','五','六']
                    for i, p in enumerate(per_d):
                        pf = ''
                        if len(per_d) > 1 or org_d:
                            idx = i + len(org_d)
                            pf = f'被告{nums[idx] if idx < len(nums) else str(idx+1)} '
                        _label_data_row(table, label, _fmt_person(p, pf))
                else:
                    _label_data_row(table, label, _empty_person())
            elif '被告（法人' in label:
                org_d = data_context['org_defendants']
                per_d = data_context['per_defendants']
                if org_d:
                    nums = ['一','二','三','四','五','六']
                    for i, p in enumerate(org_d):
                        pf = ''
                        if len(org_d) > 1 or per_d:
                            pf = f'被告{nums[i] if i < len(nums) else str(i+1)} '
                        _label_data_row(table, label, _fmt_org(p, pf))
                else:
                    _label_data_row(table, label, _empty_org())
            elif '被告' in label:
                # Generic 被告 row (离婚纠纷 uses this)
                per_d = data_context['per_defendants']
                if per_d:
                    _label_data_row(table, label, _fmt_person(per_d[0]))
                else:
                    _label_data_row(table, label, _empty_person())
            elif '第三人（自然人）' in label:
                per_t = data_context['per_thirds']
                if per_t:
                    _label_data_row(table, label, _fmt_person(per_t[0]))
                else:
                    _label_data_row(table, label, _empty_person())
            elif '第三人（法人' in label and '类型' not in label:
                org_t = data_context['org_thirds']
                if org_t:
                    _label_data_row(table, label, _fmt_org(org_t[0]))
                else:
                    _label_data_row(table, label, _empty_org())
            elif '委托诉讼代理人' in label:
                agent = data_context['agent']
                if agent.get('has') or agent.get('name'):
                    _label_data_row(table, label, _fmt_agent(agent))
                else:
                    _label_data_row(table, label, _empty_agent())
            
            # ── Mediation rows ──
            elif '是否了解先行调解解决纠纷的好处' in label:
                mediation_benefit_count += 1
                if mediation_benefit_count == 1:
                    # Benefits 1-4
                    _label_data_row(table, label, [
                        '1.立案后选择先行调解的，可以很快启动调解程序。如不同意调解，法院将依程序开庭审理案件，但可能需要经过较长一段时间才能拿到判决结果。',
                        '2.先行调解不收取任何费用，如果后续达成调解协议，亦无需交纳案件受理费。',
                        '3.先行调解可以选择您认为方便的方式，包括线上调解或线下调解。',
                        '4.与开庭审理相比，先行调解的程序更加灵活、便捷。',
                    ])
                else:
                    # Benefit 5
                    _label_data_row(table, label, [
                        '5.调解达成的协议具有法律效力，可以依照法律规定申请司法确认，具有强制执行效力。',
                        '了解□   不了解□',
                    ])
            
            else:
                # Generic data-matching row
                data_lines = _match_label_to_data(label, case_data)
                _label_data_row(table, label, data_lines)
    
    return table


def generate_docx(data, output_path):
    """Main entry point: generate DOCX from case data."""
    _load_data()
    
    case_type_name = data.get('caseTypeName', '民事纠纷')
    
    # ── Find matching layout ──
    layout = None
    # Try exact match
    if case_type_name in _table_layouts:
        layout = _table_layouts[case_type_name]
    else:
        # Try fuzzy match
        for name, l in _table_layouts.items():
            if case_type_name in name or name in case_type_name:
                layout = l
                break
        # Fallback: try TYPE_ID_TO_NAME
        if not layout:
            case_id = data.get('caseTypeId', '')
            names = _templates.get('TYPE_ID_TO_NAME', {})
            mapped_name = names.get(case_id, '')
            if mapped_name and mapped_name in _table_layouts:
                layout = _table_layouts[mapped_name]
    
    if not layout:
        # Fallback: use JRJK layout as default
        layout = _table_layouts.get('金融借款合同纠纷')
        if not layout and _table_layouts:
            layout = list(_table_layouts.values())[0]
        if not layout:
            raise ValueError(f'No layout found for {case_type_name} and no fallback available')
    
    # ── Prepare data context ──
    plaintiffs = data.get('plaintiffs', [])
    defendants = data.get('defendants', [])
    thirds = data.get('thirds', [])
    agent = data.get('agent', {})
    
    ctx = {
        'per_plaintiffs': [p for p in plaintiffs if p.get('type') != 'org'],
        'org_plaintiffs': [p for p in plaintiffs if p.get('type') == 'org'],
        'per_defendants': [d for d in defendants if d.get('type') != 'org'],
        'org_defendants': [d for d in defendants if d.get('type') == 'org'],
        'per_thirds': [t for t in thirds if t.get('type') != 'org'],
        'org_thirds': [t for t in thirds if t.get('type') == 'org'],
        'agent': agent,
    }
    
    # ── Build document ──
    doc = Document()
    for sec in doc.sections:
        sec.page_width  = Emu(PAGE_W * 635)
        sec.page_height = Emu(PAGE_H * 635)
        sec.top_margin    = Emu(MARGIN_TOP)
        sec.bottom_margin = Emu(MARGIN_BOTTOM)
        sec.left_margin   = Emu(MARGIN_LEFT)
        sec.right_margin  = Emu(MARGIN_RIGHT)
    
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = SZ
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    
    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _run(p, '民事起诉状', bold=True, size=SZ_TTL)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _run(p, f'（{case_type_name}）', size=SZ_SUB)
    
    # ── Render tables from layout ──
    for tbl_layout in layout['tables']:
        _render_table(doc, tbl_layout, data, ctx)
    
    # ── Signature ──
    juris = data.get('jurisdiction', {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(0)
    _run(p, '具状人（签字、盖章）：')
    
    date_str = juris.get('date', '') or datetime.now().strftime('%Y.%m.%d')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _run(p, f'日期：{date_str}')
    
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: generate_docx.py <data.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)
    result = generate_docx(data, sys.argv[2])
    print(f'DOCX generated: {result}')
