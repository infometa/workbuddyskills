#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MAI Cap-Table 勾稽校验闸门 (Reconciliation Gate) — Phase 1b
============================================================
专抓 grounding_gate 抓不住的"来源对、算术错"持股表错误:
  典型问题一: 董事权益栏(个人+受控法团合计) 与 主要股东表(法团单列) 是同一笔股份两种口径,
          相加致实控比例翻倍 (控制人甲 31.15% + 受控法团A 30.95% = 错误 62.1%)。
          机器信号: 两行绝对股数近乎相等 -> 疑为同一笔双重披露。
  典型问题二: SFO s.336 主要股东表各行 % 可能基于不同申报日的股本基数, 加总前须统一折算。
          机器信号: 各行"隐含股本基数 = 绝对股数 / 比例"不一致。

设计要点 (基于脱敏样本回测得出):
  - 必须先"按表头识别单张持股表"(同时含 股数列 与 比例列)再校验; 直接整文件扫会把
    多表 Excel 的不同表混读 -> 灾难性误报(v1 教训)。用表头定位列, 遇下一表头/空行/sheet界结束。
  - 分母一致性检查只看重大持股行(比例≥1%): 小比例的2位四舍五入会让反算基数有±8%噪声。
  - 负值/轧差行(其余主体=总数-已披露, 常为负)剔除, 不当正常持股。

适用层: 内部测算/草稿里的持股表 (Excel/Word/文本)。
用法: recon_gate.py <file.xlsx|docx|csv|txt> [更多文件...]
退出码: 0=完成且无问题; 1=发现勾稽问题; 2=未识别持股表或文件未完成校验

回测基线 (2026-06-28):
  错误口径样例: 控制人/受控法团股数近乎相等 + 个别行隐含股本基数离群 均命中
  正确版(年报核实31.15%) / 脱敏最终测算表: 零误报
"""
import csv, re, sys, os


class UnverifiedError(Exception):
    """The requested check could not be completed."""

YI = 1e8
SHARES_HDR = re.compile(r'持股|股数|股份数目|亿股|持有股份')
PCT_HDR    = re.compile(r'占已发行|百分比|比例|持股\s*%|%')
TOTAL_ROW  = re.compile(r'合计|总计|小计|全平台|校验|公众及独立|公众股东')

def extract_lines(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md"):
        return open(path, encoding="utf-8").read().splitlines()
    if ext == ".csv":
        with open(path, encoding="utf-8-sig", newline="") as stream:
            return [" | ".join(cell.strip() for cell in row) for row in csv.reader(stream)]
    if ext == ".xlsx":
        try:
            import openpyxl
        except ImportError as exc:
            raise UnverifiedError("读取 .xlsx 需要安装 openpyxl") from exc
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
        except Exception as exc:
            raise UnverifiedError(f"无法读取 XLSX: {exc}") from exc
        out = []
        for ws in wb.worksheets:
            out.append("##### SHEET")
            for r in ws.iter_rows(values_only=True):
                cells = [("" if c is None else str(c)) for c in r]
                if any(c.strip() for c in cells): out.append(" | ".join(cells))
        return out
    if ext == ".docx":
        try:
            import docx
        except ImportError as exc:
            raise UnverifiedError("读取 .docx 需要安装 python-docx") from exc
        try:
            d = docx.Document(path)
        except Exception as exc:
            raise UnverifiedError(f"无法读取 DOCX: {exc}") from exc
        out = [p.text for p in d.paragraphs]
        for t in d.tables:
            for r in t.rows:
                out.append(" | ".join(c.text.strip() for c in r.cells))
        return out
    raise UnverifiedError(
        f"不支持的文件类型: {ext or '[无扩展名]'} (支持 .xlsx/.docx/.csv/.txt/.md)"
    )

def num(cell):
    m = re.search(r'-?[\d,]+(?:\.\d+)?', cell)   # 捕获负号: 余额/轧差行常为负, 须能识别并剔除
    if not m: return None
    try: return float(m.group().replace(",", ""))
    except: return None

def find_tables(lines):
    tables = []; i = 0
    while i < len(lines):
        cells = [c.strip() for c in lines[i].split("|")]
        sc = pc = None
        for j, c in enumerate(cells):
            if sc is None and SHARES_HDR.search(c) and "股" in c: sc = j
            if pc is None and PCT_HDR.search(c) and not SHARES_HDR.search(c): pc = j
        if sc is not None and pc is not None and sc != pc:
            unit = YI if "亿" in cells[sc] else 1
            tbl = {
                "shares_col": sc,
                "pct_col": pc,
                "unit": unit,
                "hdr_line": i + 1,
                "rows": [],
                "totals": [],
                "incomplete": [],
            }
            i += 1
            while i < len(lines):
                ln = lines[i]
                if not ln.strip() or ln.startswith("#####"): break
                cc = [c.strip() for c in ln.split("|")]
                if any(SHARES_HDR.search(c) and "股" in c for c in cc) and any(PCT_HDR.search(c) for c in cc):
                    break
                if len(cc) <= max(sc, pc): i += 1; continue
                name = cc[0]
                if not name or not re.search(r'[一-鿿A-Za-z]', name): i += 1; continue
                sval = num(cc[sc]); pval = num(cc[pc])
                if sval is None or pval is None:
                    tbl["incomplete"].append((i + 1, name[:18], cc[sc], cc[pc]))
                    i += 1
                    continue
                shares = sval * tbl["unit"]
                pct = pval if pval < 1 else pval / 100
                if shares <= 0 or pct <= 0: i += 1; continue
                rec = (i+1, name[:18], shares, pct)
                (tbl["totals"] if TOTAL_ROW.search(name) else tbl["rows"]).append(rec)
                i += 1
            if len(tbl["rows"]) >= 2 or tbl["incomplete"]: tables.append(tbl)
            continue
        i += 1
    return tables

def check_table(tbl, idx):
    rows = tbl["rows"]; issues = 0
    print(f"\n  -- 表#{idx} (表头行 L{tbl['hdr_line']}, 股数列{tbl['shares_col']}/比例列{tbl['pct_col']}, {len(rows)}行) --")
    for _, n, s, p in rows:
        print(f"     {n:<16} 股数={s:>16,.0f}  比例={p*100:6.2f}%  隐含基数={s/p/YI:8.3f}亿")
    for line, name, shares, pct in tbl["incomplete"]:
        print(f"     [UNVERIFIED] L{line} {name}: 数据不完整 (股数={shares or '[空]'}, 比例={pct or '[空]'})")
    material = [(ln, n, s, p) for ln, n, s, p in rows if p >= 0.01]
    bases = [s/p for _, _, s, p in material]
    if len(bases) >= 2:
        bmin, bmax = min(bases), max(bases)
        if bmax/bmin - 1 > 0.03:
            issues += 1
            print(f"     [WARN] 分母不一致: 隐含股本基数不一致 ({bmin/YI:.2f}~{bmax/YI:.2f}亿, 差{(bmax/bmin-1)*100:.1f}%) -> 各行%分母不同, 加总须先折算")
            med = sorted(bases)[len(bases)//2]
            for (_, n, s, p) in material:
                if abs((s/p)/med - 1) > 0.03:
                    print(f"          离群: {n} 隐含基数{s/p/YI:.2f}亿 (多数为{med/YI:.2f}亿)")
    for a in range(len(rows)):
        for b in range(a+1, len(rows)):
            sa, sb = rows[a][2], rows[b][2]
            if rows[a][1] == rows[b][1]: continue
            if min(sa, sb) > 0 and abs(sa-sb)/max(sa, sb) < 0.02:
                issues += 1
                print(f"     [WARN] 疑似双重披露: 「{rows[a][1]}」与「{rows[b][1]}」股数近乎相等 ({sa:,.0f} vs {sb:,.0f}) -> 疑同一笔双重披露, 相加=控制权翻倍, 核对nesting")
    tot = sum(p for _, _, _, p in rows)
    if tot > 1.02:
        issues += 1
        print(f"     [WARN] 比例合计 {tot*100:.1f}% >100% -> 不可能, 多半存在重复披露相加")
    row_shares = sum(s for _, _, s, _ in rows)
    row_pct = sum(p for _, _, _, p in rows)
    for _, name, total_shares, total_pct in tbl["totals"]:
        shares_gap = abs(total_shares - row_shares) / max(total_shares, row_shares)
        pct_gap = abs(total_pct - row_pct)
        if shares_gap > 0.005 or pct_gap > 0.005:
            issues += 1
            print(
                f"     [WARN] 总计行「{name}」与明细不一致: "
                f"股数 {total_shares:,.0f} vs {row_shares:,.0f}, "
                f"比例 {total_pct*100:.2f}% vs {row_pct*100:.2f}%"
            )
    unverified = bool(tbl["incomplete"])
    if issues == 0 and not unverified: print("     [OK] 勾稽通过")
    return issues, unverified

def check_file(path):
    tables = find_tables(extract_lines(path))
    print(f"\n### {os.path.basename(path)} — 识别出 {len(tables)} 张持股表")
    if not tables:
        print("  [UNVERIFIED] 未完成校验：未识别到含'股数列+比例列'的持股表")
        return None
    checked = [check_table(table, index + 1) for index, table in enumerate(tables)]
    return sum(item[0] for item in checked), any(item[1] for item in checked)

def main(argv=None):
    paths = [a for a in (argv if argv is not None else sys.argv[1:]) if not a.startswith("--")]
    if not paths:
        print(__doc__)
        return 2
    total = 0
    unverified = 0
    for p in paths:
        try:
            result = check_file(p)
        except (OSError, UnverifiedError, ValueError) as exc:
            result = None
            print(f"\n### {os.path.basename(p)}")
            print(f"  [UNVERIFIED] 未完成校验: {exc}")
        if result is None:
            unverified += 1
        else:
            issues, has_unverified_rows = result
            total += issues
            unverified += int(has_unverified_rows)
    if total:
        print(f"\n{'='*50}\n合计勾稽问题: {total} -> [FAIL] 须核对原文后修正")
        return 1
    if unverified:
        print(f"\n{'='*50}\n未完成校验: {unverified} 个文件 -> [UNVERIFIED] 未完成校验")
        return 2
    status = "[FAIL] 须核对年报原文后修正" if total else "[OK] 勾稽通过"
    print(f"\n{'='*50}\n合计勾稽问题: {total} -> {status}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
