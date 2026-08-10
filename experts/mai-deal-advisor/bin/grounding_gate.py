#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MAI Source-Grounding Gate — 出处分层校验闸门
============================================================
把"上市公司股东数据必须溯源至年报主要股东章"纪律
变成机器强制: 扫描内部测算/草稿里的对外数字, 判断其来源层级是否达标。

核心规则 (基于脱敏样本回测得出):
  - 典型失败不是"数字没出处", 是"出处引错层级"(broker/披露易当终结证据)。
    因此朴素"有没有引用"的检查会放行错误数字 —— 必须按"声明类型→要求来源层级"判定。
  - 高风险类型(上市公司股东持股/股本): 只有年报/招股书等一手原文算终结锚点;
    broker/HKEx披露易/财经网站即便被引用, 也判 NEEDS_PRIMARY(须一手原文确认)。
  - 表底【数据来源】里的一手锚点可继承给整张表的行 (避免逐行重复)。
  - 测算/情景输出(Step2持股/增长倍率/估值假设)是模型推导, 不在此闸门拦截;
    其算术正确性由独立的勾稽校验(Phase 1b)负责。

适用层: 内部测算/草稿 (Excel/Word/文本), 不是对外交付稿(交付稿不应内联年报页码)。

用法:
  grounding_gate.py <file.xlsx|file.docx|file.csv|file.pdf|file.txt> [更多文件...]
退出码: 0=完成且无拦截项; 1=发现应拦截项; 2=文件未完成校验

回测基线 (2026-06-28):
  修正前辅助源口径样例: 某主要股东比例 -> NEEDS_PRIMARY (成功拦截错误来源层级)
  已修正版:             同一比例 -> PASS (继承"年报第52页"锚点); 误报仅 4 条且均合理
"""
import csv, re, sys, os
from collections import Counter


class UnverifiedError(Exception):
    """The requested check could not be completed."""

# ---------- 来源层级词表 ----------
TERMINAL  = re.compile(r'年报|年度报告|中期报告|招股书|招股章程|经审计|审计报告|财务报表|收购守则|上市规则')
AUXILIARY = re.compile(r'broker|券商|富途|futu|hafoo|哈富|披露易|DION|DI数据库|HKEx|HKEX|wikipedia|维基|新浪|东方财富|百度|雪球|同花顺')
ASSUMPTION= re.compile(r'假设|测算|拟议|协商确定|客户口径|客户最新口径|目标值|目标比例|目标为|情景|预计|估计|约定|建议值|锁定值')
DEAL_DOC  = re.compile(r'公告|备忘录|MOU|Term\s*Sheet|协议|合同')
SOURCE_LABEL = re.compile(r'数据来源|资料来源|来源[:：]|source[:：]|注[:：]', re.I)
LOCATOR = re.compile(
    r'https?://\S+|第\s*\d+\s*页|\b(?:p\.?|page)\s*\d+\b|'
    r'公告编号\s*[:：]?\s*[A-Za-z0-9._/-]+|'
    r'(?:工作表|sheet)\s*[:：]?\s*[^,，；;|]+\s*(?:单元格|cell)\s*[:：]?\s*[A-Z]+\d+',
    re.I,
)

# ---------- 区块切分 / 类型 ----------
HEADER      = re.compile(r'^\s*#{1,6}\s+|^\s*[一二三四五六七八九十]+\s*、|^\s*【|方案\s*Step|Step\s*2\s*完成后|^\s*\d+\.\s')
FACT_HDR    = re.compile(r'股东比例|主要股东|内部股东|当前状态|交易前|已发行股本|股本')
DERIVED_HDR = re.compile(r'Step\s*2\s*完成后|情景|价值与收益|持股价值|投入基数|平台参数|关键观察|数据来源|关于')

# ---------- 数字识别 ----------
PCT    = re.compile(r'\d+(?:\.\d+)?\s*%|0\.\d{2,4}\b')
SHARES = re.compile(r'[\d,]+(?:\.\d+)?\s*(?:亿|万)?\s*股')
MONEY  = re.compile(r'(?:港币|人民币|USD|US\$|HK\$|RMB)\s*[\d,.–\-]+\s*(?:亿|万)?|[\d,]+(?:\.\d+)?\s*亿(?:元|美元|港元)?')

BLOCKERS = ("NEEDS_PRIMARY", "UNANCHORED")

# ============================================================
# 文件 -> 文本行
# ============================================================
def extract_lines(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md"):
        return open(path, encoding="utf-8").read().splitlines()
    if ext == ".csv":
        with open(path, encoding="utf-8-sig", newline="") as stream:
            return [" | ".join(cell.strip() for cell in row) for row in csv.reader(stream)]
    if ext == ".xlsx":
        try:
            import openpyxl
        except ImportError as exc:
            raise UnverifiedError("读取 .xlsx 需要安装 openpyxl") from exc
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
        except Exception as exc:
            raise UnverifiedError(f"无法读取 XLSX: {exc}") from exc
        lines = []
        for ws in wb.worksheets:
            lines.append(f"##### SHEET: {ws.title}")
            for r in ws.iter_rows(values_only=True):
                cells = [str(c) for c in r if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return lines
    if ext == ".docx":
        try:
            import docx
        except ImportError as exc:
            raise UnverifiedError("读取 .docx 需要安装 python-docx") from exc
        try:
            d = docx.Document(path)
        except Exception as exc:
            raise UnverifiedError(f"无法读取 DOCX: {exc}") from exc
        lines = [p.text for p in d.paragraphs]
        for t in d.tables:
            for r in t.rows:
                cells = [c.text.strip() for c in r.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return lines
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise UnverifiedError("读取 .pdf 需要安装 pypdf") from exc
        try:
            reader = PdfReader(path)
            lines = []
            for page in reader.pages:
                lines.extend((page.extract_text() or "").splitlines())
        except Exception as exc:
            raise UnverifiedError(f"无法提取 PDF 文本: {exc}") from exc
        if not any(line.strip() for line in lines):
            raise UnverifiedError("PDF 没有可提取文本；扫描件需要 OCR 后人工复核")
        return lines
    raise UnverifiedError(
        f"不支持的文件类型: {ext or '[无扩展名]'} "
        "(支持 .xlsx/.docx/.csv/.pdf/.txt/.md)"
    )

# ============================================================
# 核心判定
# ============================================================
def find_numbers(text):
    spans = []
    for pat in (PCT, SHARES, MONEY):
        for m in pat.finditer(text):
            spans.append((m.start(), m.end(), m.group().strip()))
    spans.sort(); out = []
    for s, e, g in spans:
        if out and s < out[-1][1]:
            if e - s > out[-1][1] - out[-1][0]: out[-1] = (s, e, g)
            continue
        out.append((s, e, g))
    return [(s, e, g) for s, e, g in out if g and g not in {"%", "股", "亿"}]

def claim_type(ctx):
    if re.search(r'持股|股东|占已发行|实控|实际控制|股权|表决权|持\s*[\d,]', ctx): return "SHAREHOLDING"
    if re.search(r'已发行股本|总股本|股本', ctx): return "SHARE_CAPITAL"
    if re.search(r'估值|对价|收购|认购|市值|出资|投入|募集|并购贷款', ctx): return "VALUATION"
    if re.search(r'收入|营收|利润|EBITDA|现金流|总资产|净资产|负债|毛利|净利', ctx, re.I): return "FINANCIAL"
    return "INFO"

def segment(lines):
    blocks = []; cur = {"lines": []}
    for i, ln in enumerate(lines):
        if (HEADER.match(ln) or ln.startswith("##### SHEET")) and cur["lines"]:
            blocks.append(cur); cur = {"lines": []}
        cur["lines"].append((i, ln))
    if cur["lines"]: blocks.append(cur)
    m = {}
    for b in blocks:
        text = " ".join(l for _, l in b["lines"])
        b["has_source_anchor"] = any(
            TERMINAL.search(line) and SOURCE_LABEL.search(line) and LOCATOR.search(line)
            for _, line in b["lines"]
        )
        first = b["lines"][0][1]
        b["kind"] = "DERIVED" if (DERIVED_HDR.search(first) and not FACT_HDR.search(first)) else "FACT"
        for i, _ in b["lines"]: m[i] = b
    return m

def verdict(line, ctype, block):
    has_term = bool(TERMINAL.search(line) and LOCATOR.search(line)); has_aux = bool(AUXILIARY.search(line))
    has_asm  = bool(ASSUMPTION.search(line)); has_deal = bool(DEAL_DOC.search(line))
    has_deal = bool(has_deal and LOCATOR.search(line))
    inherit  = bool(block and block.get("has_source_anchor"))
    derived  = bool(block and block["kind"] == "DERIVED")
    if ctype in ("SHAREHOLDING", "SHARE_CAPITAL"):
        if has_asm or derived: return "OK_DERIVED", "已明确标注的测算/推导输出(正确性走勾稽校验)"
        if has_term:  return "PASS", "行内一手锚点"
        if inherit:   return "PASS_INHERIT", "继承表级一手锚点(数据来源脚注)"
        if has_aux:   return "NEEDS_PRIMARY", "仅辅助源(broker/披露易等)→须一手原文(年报主要股东章)确认"
        return "UNANCHORED", "无任何出处→待确认"
    if ctype in ("VALUATION", "FINANCIAL"):
        if has_term or has_deal: return "PASS", "一手/交易文件锚点"
        if has_asm or derived:   return "OK_ASSUMPTION", "已标假设/测算口径"
        return "UNANCHORED", "财务/估值数字无来源且未标假设→待确认"
    return "INFO", "信息性数字(非高风险)"

def run(path):
    lines = extract_lines(path)
    bmap = segment(lines)
    led = []
    for i, ln in enumerate(lines):
        if not ln.strip() or ln.startswith("#####"): continue
        nums = find_numbers(ln)
        if not nums: continue
        ct = claim_type(ln); v, r = verdict(ln, ct, bmap.get(i))
        for s, e, g in nums:
            led.append({"line": i + 1, "num": g, "ctype": ct, "verdict": v, "reason": r, "ctx": ln.strip()[:95]})
    return led

def report(path):
    led = run(path)
    covered = [item for item in led if item["ctype"] != "INFO"]
    if not covered:
        raise UnverifiedError("未识别到可检查的持股、股本、估值或财务数字")
    vc = Counter(x["verdict"] for x in led)
    block = [x for x in led if x["verdict"] in BLOCKERS]
    print(f"\n### {os.path.basename(path)}  (数字 claims={len(led)})")
    for v in ("PASS", "PASS_INHERIT", "OK_DERIVED", "OK_ASSUMPTION", "NEEDS_PRIMARY", "UNANCHORED", "INFO"):
        if vc.get(v): print(f"  {v:<14}: {vc[v]}")
    print(f">> 应拦截(须补一手原文/待确认): {len(block)} 条")
    seen = set()
    for x in block:
        k = (x["num"], x["ctx"][:25])
        if k in seen: continue
        seen.add(k)
        print(f"   L{x['line']:>3} [{x['verdict']:<13}] {x['num']:<12}| {x['ctx'][:64]}")
    return len(block)

def main(argv=None):
    paths = [a for a in (argv if argv is not None else sys.argv[1:]) if not a.startswith("--")]
    if not paths:
        print(__doc__)
        return 2
    total_block = 0
    unverified = 0
    for p in paths:
        try:
            total_block += report(p)
        except (OSError, UnverifiedError, ValueError) as exc:
            unverified += 1
            print(f"\n### {os.path.basename(p)}")
            print(f">> [UNVERIFIED] 未完成校验: {exc}")
    if unverified:
        print(f"\n{'='*50}\n未完成校验: {unverified} 个文件 -> [UNVERIFIED] 未完成校验")
        return 2
    status = "[FAIL] 不可交付,先补一手锚点" if total_block else "[OK] 出处层级达标"
    print(f"\n{'='*50}\n合计应拦截: {total_block} 条 -> {status}")
    return 1 if total_block else 0


if __name__ == "__main__":
    sys.exit(main())
