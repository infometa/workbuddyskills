#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
患教内容审核 · Word 批注生成脚本
========================================
把待审文章（.docx）原文保留，对存在问题的句子/词组做高亮，并挂上 Word 原生批注
（右侧批注气泡，可在 Word / WPS 中直接查看、回复、接受修改），文末追加结构化审核结论页。

用法：
  python review_to_word.py --src 待审文章.docx --issues issues.json --out 已审核_带批注.docx

issues.json 结构（由审核专家依据六维度审核结果生成）：
{
  "title": "文章标题（可选，用于结论页）",
  "summary": [                     # 审核结论表，每维度一行（result 内不要放彩色圆点 emoji）
    {"dimension": "严重问题检查", "result": "未发现问题"},
    {"dimension": "科学性审核",   "result": "建议优化（2 项）"},
    ...
  ],
  "issues": [
    {
      "anchor": "格外接地气",       # 在原文中要定位/高亮的文本片段（尽量短、唯一）
      "level": "建议优化",          # 需修改 / 建议优化 / 可接受（也兼容 🔴/🟡/🟢，脚本会自动归一化为文字）
      "dimension": "出版规范",       # 所属审核维度
      "problem": "编辑口吻/主观评价", # 问题描述
      "suggestion": "建议删除，直接陈述主题" # 修改/优化建议
    },
    ...
  ]
}

依赖：python-docx >= 1.1（需支持 Document.add_comment）
"""
import argparse, json, sys, os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

# 分级 → 高亮色（同时兼容 emoji 键与纯文字键，避免 Word/WPS 无法显示彩色圆点的问题）
LEVEL_COLOR = {
    "🔴": WD_COLOR_INDEX.RED,
    "🟡": WD_COLOR_INDEX.YELLOW,
    "🟢": WD_COLOR_INDEX.BRIGHT_GREEN,
    "需修改": WD_COLOR_INDEX.RED,
    "建议优化": WD_COLOR_INDEX.YELLOW,
    "可接受": WD_COLOR_INDEX.BRIGHT_GREEN,
    "备注": WD_COLOR_INDEX.BRIGHT_GREEN,
}
# 分级 → 写入 Word 的纯文字标签（不用 emoji，任何环境都能正常显示）
LEVEL_NAME = {
    "🔴": "需修改", "🟡": "建议优化", "🟢": "可接受",
    "需修改": "需修改", "建议优化": "建议优化", "可接受": "可接受", "备注": "可接受",
}
# 归一化：把各种写法映射到标准文字标签
LEVEL_CANON = {
    "🔴": "需修改", "🟡": "建议优化", "🟢": "可接受",
    "需修改": "需修改", "建议优化": "建议优化", "可接受": "可接受", "备注": "可接受",
}
# 分级 → 批注气泡前缀 emoji（批注气泡为纯文本，emoji 可正常显示；仅正文高亮/结论表忌用 emoji）
LEVEL_EMOJI = {
    "🔴": "🔴", "🟡": "🟡", "🟢": "🟢",
    "需修改": "🔴", "建议优化": "🟡", "可接受": "🟢", "备注": "🟢",
}
AUTHOR = "药箱审核助手"
INITIALS = "审"


def find_paragraph_with(doc, anchor):
    """返回第一个包含 anchor 文本的段落对象，找不到返回 None。"""
    for p in doc.paragraphs:
        if anchor and anchor in p.text:
            return p
    # 也检查表格内段落
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if anchor and anchor in p.text:
                        return p
    return None


def highlight_and_comment(doc, p, anchor, level, comment_text):
    """在段落 p 中定位 anchor，仅高亮该片段并挂批注。"""
    full = p.text
    idx = full.find(anchor)
    if idx < 0:
        return False
    # 保留原段落样式：记录首个 run 的字体基本属性
    base_font = None
    if p.runs:
        base_font = p.runs[0].font
    # 清空原 runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    pre, mid, post = full[:idx], full[idx:idx + len(anchor)], full[idx + len(anchor):]

    def _mkrun(text):
        run = p.add_run(text)
        # 尽量继承原字号（避免变形）
        if base_font is not None and base_font.size:
            run.font.size = base_font.size
        return run

    if pre:
        _mkrun(pre)
    run_mid = _mkrun(mid)
    run_mid.font.highlight_color = LEVEL_COLOR.get(level, WD_COLOR_INDEX.YELLOW)
    if post:
        _mkrun(post)
    doc.add_comment(runs=[run_mid], text=comment_text, author=AUTHOR, initials=INITIALS)
    return True


def _strip_emoji_dots(text):
    """把结论文字里的彩色圆点 emoji 去掉，避免 Word/WPS 显示成方框。"""
    if not text:
        return text
    for e in ("🔴", "🟡", "🟢", "🔺", "🔻"):
        text = text.replace(e, "")
    return text.strip()


def _set_table_style(table):
    """安全设置表格边框样式，样式缺失时回退。"""
    for name in ("Table Grid", "TableGrid"):
        try:
            table.style = name
            return
        except Exception:
            continue
    # 最后兜底：不设样式（仍能显示，只是无边框）


def add_summary_pages(doc, data):
    """文末追加审核结论页 + 明细清单（未能在原文定位的问题也在此兜底列出）。"""
    doc.add_page_break()
    h = doc.add_heading("患教内容审核报告", level=0)

    if data.get("title"):
        doc.add_paragraph("审核对象：%s" % data["title"])

    # 结论表
    doc.add_heading("一、审核结论", level=1)
    summary = data.get("summary", [])
    if summary:
        table = doc.add_table(rows=1, cols=2)
        _set_table_style(table)
        hdr = table.rows[0].cells
        hdr[0].text = "审核维度"
        hdr[1].text = "结果"
        for row in summary:
            cells = table.add_row().cells
            cells[0].text = _strip_emoji_dots(str(row.get("dimension", "")))
            cells[1].text = _strip_emoji_dots(str(row.get("result", "")))

    # 明细清单（按 需修改 / 建议优化 / 可接受 分组；兼容 issues 里 level 用 emoji 或文字）
    doc.add_heading("二、详细审核意见", level=1)
    issues = data.get("issues", [])
    for canon in ["需修改", "建议优化", "可接受"]:
        group = [it for it in issues if LEVEL_CANON.get(it.get("level"), "") == canon]
        if not group:
            continue
        doc.add_heading("%s（%d 项）" % (canon, len(group)), level=2)
        table = doc.add_table(rows=1, cols=4)
        _set_table_style(table)
        hdr = table.rows[0].cells
        for i, t in enumerate(["维度", "位置/原文", "问题描述", "修改建议"]):
            hdr[i].text = t
        for it in group:
            cells = table.add_row().cells
            cells[0].text = str(it.get("dimension", ""))
            cells[1].text = str(it.get("anchor", it.get("location", "")))
            cells[2].text = str(it.get("problem", ""))
            cells[3].text = str(it.get("suggestion", ""))

    note = doc.add_paragraph()
    note.add_run(
        "说明：正文中带底色高亮的文字为问题定位，可在批注栏查看对应审核意见"
        "（红色高亮=需修改，黄色高亮=建议优化，绿色高亮=可接受）。"
        "本报告由 AI 审核助手生成，仅供内容合规预审参考，最终发布决策请以人工复核为准。"
    )
    for r in note.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", required=True, help="待审文章 .docx 路径")
    ap.add_argument("--issues", required=True, help="问题清单 issues.json 路径")
    ap.add_argument("--out", required=True, help="输出的带批注 .docx 路径")
    args = ap.parse_args()

    if not os.path.isfile(args.src):
        print("[ERR] 找不到待审文章:", args.src); sys.exit(1)
    with open(args.issues, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(args.src)

    matched, unmatched = 0, []
    # 注意：遍历全部 issues —— 需修改(🔴) 与 建议优化(🟡) 均须批注，不得只批 🔴 而省略 🟡。
    for it in data.get("issues", []):
        anchor = it.get("anchor", "").strip()
        level = it.get("level", "🟡")
        label = LEVEL_CANON.get(level, "建议优化")
        emoji = LEVEL_EMOJI.get(level, "🟡")
        # 批注格式：[🔴/🟡][维度][问题描述][修改建议]
        comment_text = "[%s %s]  [维度：%s]\n[问题] %s%s" % (
            emoji, label,
            it.get("dimension", ""),
            it.get("problem", ""),
            "\n[修改建议] " + it.get("suggestion", "") if it.get("suggestion") else "",
        )
        p = find_paragraph_with(doc, anchor) if anchor else None
        if p is not None and highlight_and_comment(doc, p, anchor, level, comment_text):
            matched += 1
        else:
            unmatched.append(anchor or it.get("problem", ""))

    add_summary_pages(doc, data)
    doc.save(args.out)
    print("[OK] 已生成带批注文档:", args.out)
    print("     命中并批注 %d 项；未在原文定位 %d 项（已在报告明细中列出）" % (matched, len(unmatched)))
    if unmatched:
        print("     未定位锚点:", " | ".join(x[:20] for x in unmatched))


if __name__ == "__main__":
    main()
